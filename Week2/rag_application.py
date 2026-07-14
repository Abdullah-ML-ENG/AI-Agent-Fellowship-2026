import os
import json
import uuid
import shutil
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional

import streamlit as st
from dotenv import load_dotenv

from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_community.vectorstores import FAISS
from pypdf import PdfReader
from chromadb.config import Settings

# ---------------------------
# Config & Setup
# ---------------------------
load_dotenv()

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=LOG_LEVEL, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger("rag_app")

DATA_DIR = Path(os.getenv("DATA_DIR", "./data"))
RAW_DOCS_DIR = Path(os.getenv("RAW_DOCS_DIR", "./data/raw_docs"))
CHROMA_DIR = Path(os.getenv("CHROMA_DIR", "./data/chroma"))
FAISS_DIR = Path(os.getenv("FAISS_DIR", "./data/faiss"))
SESSIONS_DB = Path(os.getenv("SESSIONS_DB", "./data/sessions/sessions.json"))
DOC_REGISTRY = Path(os.getenv("DOC_REGISTRY", "./data/doc_registry.json"))

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-mpnet-base-v2")
TOP_K = int(os.getenv("TOP_K", "5"))
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "900"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))

for p in [DATA_DIR, RAW_DOCS_DIR, CHROMA_DIR, FAISS_DIR, SESSIONS_DB.parent]:
    p.mkdir(parents=True, exist_ok=True)

if not DOC_REGISTRY.exists():
    DOC_REGISTRY.write_text(json.dumps({"documents": []}, indent=2))

if not SESSIONS_DB.exists():
    SESSIONS_DB.write_text(json.dumps({"sessions": {}}, indent=2))

# ---------------------------
# Helpers
# ---------------------------
def load_registry() -> Dict:
    return json.loads(DOC_REGISTRY.read_text())

def save_registry(data: Dict):
    DOC_REGISTRY.write_text(json.dumps(data, indent=2))

def load_sessions() -> Dict:
    return json.loads(SESSIONS_DB.read_text())

def save_sessions(data: Dict):
    SESSIONS_DB.write_text(json.dumps(data, indent=2))

def clean_text(text: str) -> str:
    if not text:
        return ""
    return " ".join(text.replace("\x00", " ").split())

def file_size_human(size_bytes: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    i = 0
    s = float(size_bytes)
    while s >= 1024 and i < len(units) - 1:
        s /= 1024
        i += 1
    return f"{s:.2f} {units[i]}"

@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

def get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

def extract_pdf(file_path: Path) -> List[Document]:
    docs = []
    reader = PdfReader(str(file_path))
    for i, page in enumerate(reader.pages):
        text = clean_text(page.extract_text() or "")
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "page": i + 1,
                        "doc_id": file_path.stem,
                        "file_path": str(file_path),
                    },
                )
            )
    return docs

def extract_txt_md(file_path: Path) -> List[Document]:
    text = clean_text(file_path.read_text(encoding="utf-8", errors="ignore"))
    if not text:
        return []
    return [Document(page_content=text, metadata={
        "source": file_path.name,
        "page": None,
        "doc_id": file_path.stem,
        "file_path": str(file_path),
    })]

def extract_docx(file_path: Path) -> List[Document]:
    try:
        import docx
        d = docx.Document(str(file_path))
        parts = []

        for p in d.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        for table in d.tables:
            for row in table.rows:
                row_text = " | ".join((cell.text or "").strip() for cell in row.cells).strip(" |")
                if row_text:
                    parts.append(row_text)

        text = clean_text("\n".join(parts))
        if not text:
            return []

        return [Document(page_content=text, metadata={
            "source": file_path.name,
            "page": None,
            "doc_id": file_path.stem,
            "file_path": str(file_path),
        })]
    except Exception as e:
        logger.exception(f"DOCX extraction failed for {file_path.name}: {e}")
        return []

def process_uploaded_file(uploaded_file) -> Dict:
    file_id = str(uuid.uuid4())
    ext = Path(uploaded_file.name).suffix.lower()
    local_path = RAW_DOCS_DIR / f"{file_id}_{uploaded_file.name}"
    with open(local_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if ext == ".pdf":
        raw_docs = extract_pdf(local_path)
        pages = len(raw_docs)
    elif ext in [".txt", ".md"]:
        raw_docs = extract_txt_md(local_path)
        pages = 1
    elif ext == ".docx":
        raw_docs = extract_docx(local_path)
        pages = 1
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    splitter = get_splitter()
    chunks = splitter.split_documents(raw_docs)

    if not chunks:
        raise ValueError(
            f"No extractable text/chunks found in {uploaded_file.name}. "
            "The file may be empty, image-only, or unsupported formatting."
        )

    for idx, c in enumerate(chunks):
        c.metadata["chunk_id"] = idx
        c.metadata["file_id"] = file_id

    return {
        "file_id": file_id,
        "file_name": uploaded_file.name,
        "file_path": str(local_path),
        "file_size": uploaded_file.size,
        "pages": pages,
        "chunks_count": len(chunks),
        "raw_docs": raw_docs,
        "chunks": chunks,
        "uploaded_at": datetime.utcnow().isoformat(),
        "status": "processed",
    }

def _load_or_create_faiss(embeddings):
    index_path = FAISS_DIR / "index"
    if index_path.exists():
        return FAISS.load_local(str(index_path), embeddings, allow_dangerous_deserialization=True)
    return None

def _get_chroma(embeddings):
    return Chroma(
        collection_name="enterprise_docs",
        embedding_function=embeddings,
        persist_directory=str(CHROMA_DIR),
        client_settings=Settings(is_persistent=True, anonymized_telemetry=False),
    )

def index_chunks(chunks: List[Document]):
    if not chunks:
        raise ValueError("No chunks to index.")

    embeddings = get_embeddings()

    chroma = _get_chroma(embeddings)
    chroma.add_documents(chunks)

    faiss_db = _load_or_create_faiss(embeddings)
    if faiss_db is None:
        faiss_db = FAISS.from_documents(chunks, embeddings)
    else:
        faiss_db.add_documents(chunks)
    faiss_db.save_local(str(FAISS_DIR / "index"))

def get_vectorstore(store_name: str):
    embeddings = get_embeddings()
    if store_name == "ChromaDB":
        return _get_chroma(embeddings)
    if store_name == "FAISS":
        index_path = FAISS_DIR / "index"
        if not index_path.exists():
            return None
        return FAISS.load_local(str(index_path), embeddings, allow_dangerous_deserialization=True)
    return None

def retrieve_context(question: str, store_name: str, top_k: int, filter_source: Optional[str] = None):
    vs = get_vectorstore(store_name)
    if vs is None:
        return []

    if filter_source:
        docs = vs.similarity_search(question, k=top_k, filter={"source": filter_source}) \
            if store_name == "ChromaDB" else vs.similarity_search(question, k=top_k)
        if store_name == "FAISS":
            docs = [d for d in docs if d.metadata.get("source") == filter_source]
        return docs
    return vs.similarity_search(question, k=top_k)

def build_prompt(question: str, history: List[Dict], context_docs: List[Document]) -> str:
    history_block = "\n".join([f"{m['role']}: {m['content']}" for m in history[-8:]])
    context_block = "\n\n".join([
        f"[{i}] source={d.metadata.get('source')} page={d.metadata.get('page')} chunk={d.metadata.get('chunk_id')}\n{d.page_content}"
        for i, d in enumerate(context_docs, start=1)
    ])

    return f"""
You are an enterprise document intelligence assistant.
Answer ONLY using retrieved context. If uncertain, say so explicitly.

Conversation History:
{history_block}

Retrieved Context:
{context_block}

User Question:
{question}

Requirements:
1) Provide concise answer.
2) Add citations inline in this format: [source: <doc>, page: <page>, chunk: <id>]
3) Do not invent sources.
""".strip()

def generate_answer(question: str, history: List[Dict], context_docs: List[Document]) -> str:
    if not context_docs:
        return "I could not find relevant context in the indexed documents."

    llm = ChatOpenAI(
        model=os.getenv("OPENAI_MODEL", "llama-3.1-70b-versatile"),
        temperature=0,
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_BASE_URL", "https://api.groq.com/openai/v1"),
    )
    prompt = build_prompt(question, history, context_docs)
    resp = llm.invoke(prompt)
    return getattr(resp, "content", str(resp))

def save_chat_session(session_id: str, messages: List[Dict]):
    db = load_sessions()
    db["sessions"][session_id] = {
        "updated_at": datetime.utcnow().isoformat(),
        "messages": messages,
    }
    save_sessions(db)

def load_chat_session(session_id: str):
    db = load_sessions()
    return db["sessions"].get(session_id, {}).get("messages", [])

def delete_document(file_id: str):
    reg = load_registry()
    docs = reg["documents"]
    target = None
    for d in docs:
        if d["file_id"] == file_id:
            target = d
            break
    if not target:
        return False

    try:
        Path(target["file_path"]).unlink(missing_ok=True)
    except Exception:
        pass

    remaining = [d for d in docs if d["file_id"] != file_id]
    reg["documents"] = remaining
    save_registry(reg)

    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    if FAISS_DIR.exists():
        shutil.rmtree(FAISS_DIR)
        FAISS_DIR.mkdir(parents=True, exist_ok=True)

    all_chunks = []
    for doc in remaining:
        path = Path(doc["file_path"])
        ext = path.suffix.lower()
        if ext == ".pdf":
            raw_docs = extract_pdf(path)
        elif ext in [".txt", ".md"]:
            raw_docs = extract_txt_md(path)
        elif ext == ".docx":
            raw_docs = extract_docx(path)
        else:
            continue
        chunks = get_splitter().split_documents(raw_docs)
        for idx, c in enumerate(chunks):
            c.metadata["chunk_id"] = idx
            c.metadata["file_id"] = doc["file_id"]
        all_chunks.extend(chunks)

    if all_chunks:
        index_chunks(all_chunks)

    return True

# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(page_title="Enterprise Document Intelligence Platform", layout="wide")
st.title("📚 Enterprise Document Intelligence Platform (RAG)")

if "messages" not in st.session_state:
    st.session_state.messages = []
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if "vector_store_choice" not in st.session_state:
    st.session_state.vector_store_choice = "ChromaDB"

if st.button("Load Saved Session"):
    st.session_state.messages = load_chat_session(st.session_state.session_id)
    st.success("Session loaded.")

col1, col2, col3 = st.columns([1.1, 1.4, 1.5])

with col1:
    st.subheader("📤 Upload Documents")
    uploaded_files = st.file_uploader(
        "Upload PDF/TXT/MD (DOCX bonus)",
        type=["pdf", "txt", "md", "docx"],
        accept_multiple_files=True,
    )
    if st.button("Process Uploaded Files", type="primary"):
        if not uploaded_files:
            st.warning("No files selected.")
        else:
            reg = load_registry()
            total_new_chunks = 0
            with st.spinner("Processing and indexing documents..."):
                for uf in uploaded_files:
                    try:
                        result = process_uploaded_file(uf)
                        index_chunks(result["chunks"])
                        total_new_chunks += result["chunks_count"]

                        reg["documents"].append({
                            "file_id": result["file_id"],
                            "file_name": result["file_name"],
                            "file_path": result["file_path"],
                            "file_size": result["file_size"],
                            "pages": result["pages"],
                            "chunks_count": result["chunks_count"],
                            "uploaded_at": result["uploaded_at"],
                            "status": result["status"],
                        })
                    except Exception as e:
                        logger.exception("Processing failed")
                        st.error(f"{uf.name}: processing failed -> {e}")

                save_registry(reg)
            st.success(f"Processing complete. Added chunks: {total_new_chunks}")

    st.divider()
    st.subheader("⚙️ Retrieval Settings")
    st.session_state.vector_store_choice = st.selectbox("Vector Store", ["ChromaDB", "FAISS"], index=0)
    top_k = st.slider("Top-K", min_value=1, max_value=10, value=TOP_K)
    reg = load_registry()
    doc_names = ["All Documents"] + [d["file_name"] for d in reg["documents"]]
    filter_doc = st.selectbox("Filter by Document (Advanced Feature)", doc_names)

with col2:
    st.subheader("🗂️ Document Library")
    reg = load_registry()
    docs = reg["documents"]

    total_size = sum([d["file_size"] for d in docs]) if docs else 0
    total_chunks = sum([d["chunks_count"] for d in docs]) if docs else 0

    k1, k2, k3 = st.columns(3)
    k1.metric("Total Docs", len(docs))
    k2.metric("Total Chunks", total_chunks)
    k3.metric("Storage", file_size_human(total_size))

    if not docs:
        st.info("No documents uploaded yet.")
    else:
        for d in docs:
            with st.expander(f"{d['file_name']} | {d['status']}"):
                st.write(f"**Upload Date:** {d['uploaded_at']}")
                st.write(f"**Pages:** {d['pages']}")
                st.write(f"**Chunks:** {d['chunks_count']}")
                st.write(f"**File Size:** {file_size_human(d['file_size'])}")

                c1, c2 = st.columns(2)
                if c1.button(f"Reprocess {d['file_id'][:6]}", key=f"reprocess_{d['file_id']}"):
                    try:
                        path = Path(d["file_path"])
                        ext = path.suffix.lower()
                        if ext == ".pdf":
                            raw_docs = extract_pdf(path)
                        elif ext in [".txt", ".md"]:
                            raw_docs = extract_txt_md(path)
                        elif ext == ".docx":
                            raw_docs = extract_docx(path)
                        else:
                            raise ValueError("Unsupported file type")

                        chunks = get_splitter().split_documents(raw_docs)
                        if not chunks:
                            raise ValueError("No extractable text/chunks found during reprocess.")

                        for idx, c in enumerate(chunks):
                            c.metadata["chunk_id"] = idx
                            c.metadata["file_id"] = d["file_id"]

                        index_chunks(chunks)
                        d["chunks_count"] = len(chunks)
                        d["status"] = "reprocessed"
                        save_registry(reg)
                        st.success("Reprocessed successfully.")
                    except Exception as e:
                        st.error(f"Reprocess failed: {e}")

                if c2.button(f"Delete {d['file_id'][:6]}", key=f"delete_{d['file_id']}"):
                    ok = delete_document(d["file_id"])
                    if ok:
                        st.success("Document deleted from library and vector stores.")
                        st.rerun()

with col3:
    st.subheader("💬 Chat with your documents")
    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    q = st.chat_input("Ask a question about uploaded documents...")
    if q:
        st.session_state.messages.append({"role": "user", "content": q})
        with st.chat_message("user"):
            st.markdown(q)

        selected_doc = None if filter_doc == "All Documents" else filter_doc
        retrieved = retrieve_context(
            q,
            st.session_state.vector_store_choice,
            top_k=top_k,
            filter_source=selected_doc,
        )

        with st.expander("Retrieved Chunks", expanded=False):
            if not retrieved:
                st.write("No chunks retrieved.")
            else:
                for i, r in enumerate(retrieved, start=1):
                    st.markdown(
                        f"**{i}. {r.metadata.get('source')} | page={r.metadata.get('page')} | chunk={r.metadata.get('chunk_id')}**"
                    )
                    st.write(r.page_content[:700] + ("..." if len(r.page_content) > 700 else ""))

        answer = generate_answer(q, st.session_state.messages, retrieved)

        with st.chat_message("assistant"):
            st.markdown(answer)

        st.session_state.messages.append({"role": "assistant", "content": answer})
        save_chat_session(st.session_state.session_id, st.session_state.messages)

    st.divider()
    if st.button("Export Chat (JSON)"):
        payload = {
            "session_id": st.session_state.session_id,
            "messages": st.session_state.messages,
            "exported_at": datetime.utcnow().isoformat(),
        }
        st.download_button(
            "Download Chat JSON",
            data=json.dumps(payload, indent=2),
            file_name=f"chat_{st.session_state.session_id[:8]}.json",
            mime="application/json",
        )
